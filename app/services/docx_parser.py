"""Word 文档文本提取 —— 统一处理段落 + 表格 + 页眉页脚"""

from io import BytesIO
from typing import Union


def extract_docx_text(
    source: Union[bytes, BytesIO, str],
    max_chars: int = 50000,
) -> str:
    """从 .docx 文件中提取完整文本，包括段落、表格、页眉页脚。

    Args:
        source: 文件内容 (bytes / BytesIO) 或 文件路径 (str)
        max_chars: 最大字符数，超出部分截断（0 表示不截断）

    Returns:
        提取的纯文本，段落用双换行分隔
    """
    from docx import Document

    if isinstance(source, str):
        doc = Document(source)
    elif isinstance(source, BytesIO):
        doc = Document(source)
    else:
        doc = Document(BytesIO(source))

    parts: list[str] = []

    # 1. 页面正文段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 2. 表格内容（教育类文档常用表格排版）
    for table in doc.tables:
        for row in table.rows:
            row_texts: list[str] = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                parts.append(" | ".join(row_texts))

    # 3. 页眉 / 页脚
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header:
                for para in header.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer:
                for para in footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    if not parts:
        return "[Word文件为空]"

    full_text = "\n\n".join(parts)
    if max_chars > 0 and len(full_text) > max_chars:
        full_text = full_text[:max_chars] + f"\n\n... [内容已截断，仅显示前{max_chars}字符]"

    return full_text
